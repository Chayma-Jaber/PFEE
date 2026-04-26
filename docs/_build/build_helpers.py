# -*- coding: utf-8 -*-
"""
Shared helpers for building the Barsha docx deliverables (cahier v2,
PFE report, UML annexes). Centralises styling, headings, tables, images
and TOC fields so every document looks consistent.
"""
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips
import os

PNG_DIR = os.path.join(os.path.dirname(__file__), "png")

# ---------------------------------------------------------------- styles
def set_default_styles(doc):
    """Apply consistent base styles. Body in Calibri 11, headings in Calibri."""
    styles = doc.styles
    body = styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(4)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    body.paragraph_format.line_spacing = 1.25

    for level in (1, 2, 3, 4):
        h = styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        size = {1: 18, 2: 14, 3: 12, 4: 11}[level]
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True

    # Title style
    t = styles["Title"]
    t.font.name = "Calibri"
    t.font.size = Pt(28)
    t.font.bold = True
    t.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def set_a4_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.0):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


# ---------------------------------------------------------------- content helpers
def heading(doc, text, level=1, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text, bold=False, italic=False, align=None, size=None, color=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def justified(doc, text):
    return para(doc, text, align="justify")


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(0.6 * level + 0.6)
    return p


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def code_block(doc, code, lang=""):
    """Approximate code block using a shaded paragraph with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def page_break(doc):
    doc.add_page_break()


def line_break(p):
    r = p.add_run()
    r.add_break(WD_BREAK.LINE)


def image(doc, png_name, width_cm=15.5, caption=None, fig_num=None):
    """Add an image centred, optionally with a 'Figure N — caption' label."""
    path = os.path.join(PNG_DIR, png_name)
    if not os.path.exists(path):
        para(doc, f"[image manquante: {png_name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(f"Figure {fig_num} — {caption}" if fig_num else caption)
        cr.italic = True
        cr.font.size = Pt(10)


def table(doc, headers, rows, col_widths_cm=None, header_fill="1F3A5F",
          header_color="FFFFFF", first_col_bold=False):
    """Build a styled table. Returns the table object."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"

    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        # shade cell
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_fill)
        tcPr.append(shd)

    # data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val) if val is not None else "")
            run.font.size = Pt(10)
            if first_col_bold and ci == 0:
                run.bold = True

    # column widths
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    # spacing after table
    doc.add_paragraph()
    return t


def add_toc(doc, title="Table des matières"):
    """Insert a Word TOC field. User must press F9 (or right-click → update field)
    on first open to populate it."""
    heading(doc, title, level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Cliquez ici puis appuyez sur F9 pour mettre à jour la table des matières."
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(placeholder)
    run._r.append(fldChar3)
    doc.add_page_break()


def add_field(paragraph, code, default="."):
    """Generic Word field insert (used for figure / table reference fields if needed)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = default
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, txt, fld_end):
        run._r.append(el)


def page_numbers(section):
    """Add 'Page N / M' footer to a section."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    add_field(p, "PAGE", "1")
    p.add_run(" / ")
    add_field(p, "NUMPAGES", "1")
